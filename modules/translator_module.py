import os
import requests
import time
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set environment variables
AZURE_OPENAI_ENDPOINT = os.getenv('AZURE_OPENAI_ENDPOINT')
AZURE_OPENAI_API_KEY = os.getenv('AZURE_OPENAI_API_KEY', '').strip()
DEFAULT_TARGET_LANGUAGE = os.getenv('DEFAULT_TARGET_LANGUAGE', 'pt-br')

# Check if essential variables are defined
if not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_API_KEY:
    raise ValueError("The variables AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY must be defined in the .env file")

def extract_content_from_url(url):
    """Extracts text content from paragraphs on a webpage."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        paragraphs = [p.get_text() for p in soup.find_all('p')]
        return paragraphs
    except requests.exceptions.RequestException as e:
        print(f"Error accessing the URL: {e}")
        return []

def translate_text(text, target_language=None, retries=3):
    """Translates text using the Azure OpenAI API, with retries to handle 429 errors."""
    if not target_language:
        target_language = DEFAULT_TARGET_LANGUAGE

    headers = {
        'Content-Type': 'application/json',
        'api-key': AZURE_OPENAI_API_KEY
    }

    body = {
        'messages': [
            {
                'role': 'system',
                'content': f"You are an AI assistant that translates text to {target_language}.",
            },
            {
                'role': 'user',
                'content': text,
            }
        ],
        'temperature': 0.7,
        'top_p': 0.95,
        'max_tokens': 800,
    }

    endpoint = f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/gpt-4o-mini/chat/completions?api-version=2024-08-01-preview"

    for attempt in range(retries):
        try:
            response = requests.post(endpoint, headers=headers, json=body)
            response.raise_for_status()
            response_json = response.json()

            if 'choices' in response_json and response_json['choices']:
                return response_json['choices'][0]['message']['content'].strip()
            else:
                print("Unexpected response from the API:", response_json)
                return None
        except requests.exceptions.RequestException as e:
            if response.status_code == 429:
                print("Error 429: Too many requests. Waiting to retry...")
                time.sleep(10)
            else:
                print(f"Error making request to the API: {e}")
                return None
    print("Failed to translate after several attempts.")
    return None

def translate_word_document(input_file, output_file, target_language=None):
    """Translates the content of a Word document and saves it to a new file."""
    try:
        doc = Document(input_file)
        new_doc = Document()

        for paragraph in doc.paragraphs:
            time.sleep(2)
            translated_text = translate_text(paragraph.text, target_language)
            if translated_text:
                new_doc.add_paragraph(translated_text)
                time.sleep(1)

        new_doc.save(output_file)
    except Exception as e:
        print(f"Error processing the document: {e}")
